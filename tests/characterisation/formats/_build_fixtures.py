"""Regenerate the binary reader fixtures (sample.docx, sample.pdf).

These files are committed, so this script exists for provenance rather than for
routine use: run it only when the fixtures need to change deliberately, and
review the resulting diff. The text content is kept identical to sample.md and
sample.html so the readers can be compared against one another.

    python tests/characterisation/formats/_build_fixtures.py
"""
from __future__ import annotations

import io
from pathlib import Path

HERE = Path(__file__).parent

PARAGRAPHS = [
    ("Sample document", "Heading 1"),
    ("An introductory paragraph with bold, italic and a link.", None),
    ("Requirements", "Heading 2"),
    ("You must provide evidence.", "List Bullet"),
    ("You may provide a utility bill.", "List Bullet"),
    ("A quoted passage that should not be rewritten.", "Quote"),
    ("A closing paragraph. It ends here.", None),
]


def build_docx() -> None:
    from docx import Document

    doc = Document()
    for text, style in PARAGRAPHS:
        if style and style != "Quote":
            doc.add_paragraph(text, style=style)
        else:
            doc.add_paragraph(text)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Fee"
    table.cell(1, 1).text = "250 pounds"
    doc.save(HERE / "sample.docx")


def build_pdf() -> None:
    """Write a minimal single-page PDF with no external dependency."""
    lines = [text for text, _ in PARAGRAPHS]
    content = ["BT", "/F1 12 Tf", "14 TL", "56 760 Td"]
    for line in lines:
        escaped = line.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
        content.append(f"({escaped}) Tj T*")
    content.append("ET")
    stream = "\n".join(content).encode("ascii")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /Encoding /WinAnsiEncoding >>",
    ]

    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets = []
    for number, body in enumerate(objects, start=1):
        offsets.append(out.tell())
        out.write(f"{number} 0 obj\n".encode("ascii") + body + b"\nendobj\n")

    xref_at = out.tell()
    count = len(objects) + 1
    out.write(f"xref\n0 {count}\n".encode("ascii"))
    out.write(b"0000000000 65535 f \n")
    for offset in offsets:
        out.write(f"{offset:010d} 00000 n \n".encode("ascii"))
    out.write(
        f"trailer\n<< /Size {count} /Root 1 0 R >>\nstartxref\n{xref_at}\n".encode("ascii")
    )
    out.write(b"%%EOF\n")
    (HERE / "sample.pdf").write_bytes(out.getvalue())


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print("wrote sample.docx and sample.pdf")
